import argparse, json, re
from pathlib import Path
from docx import Document
from collections import defaultdict

def norm(s: str) -> str:
    repl = {
        "−": "-", "—": "-", "–": "-", "×": "*", "∗": "*", "·": "*",
        "÷": "/", "∕": "/", "√": "sqrt", "≤": "<=", "≥": ">=", "≠": "!="
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return re.sub(r"\s+", " ", s).strip()

def looks_math(s: str) -> bool:
    return any(t in s for t in ["=", "Eq(", "Sum(", "Product(", "∑", "∀", "->", "→"])

def split_identity(s: str):
    if "=" not in s:
        return None, None
    parts = s.split("=")
    if len(parts) < 2:
        return None, None
    lhs = "=".join(parts[:-1]).strip()
    rhs = parts[-1].strip()
    if not lhs or not rhs:
        return None, None
    return lhs, rhs

def guess_symbols(lhs, rhs):
    toks = set(re.findall(r"\b([a-zA-Z]\w*)\b", lhs + " " + rhs))
    funcs = {"sin","cos","tan","log","sqrt","exp","sum","min","max","Eq","Sum","Product"}
    return [t for t in toks if t not in funcs][:12]

def extract_from_docx(path: Path, verbose=False):
    eq_lines = []
    try:
        doc = Document(str(path))
    except Exception as e:
        if verbose:
            print(f"[skip-open] {path} :: {e}")
        return eq_lines
    for p in doc.paragraphs:
        txt = norm(p.text)
        if looks_math(txt):
            eq_lines.append((txt, "para"))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                txt = norm(cell.text)
                if looks_math(txt):
                    eq_lines.append((txt, "cell"))
    if verbose:
        print(f"[scan] {path} -> {len(eq_lines)} candidate lines")
    return eq_lines

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--map", required=True)
    ap.add_argument("--verbose", action="store_true")
    args = ap.parse_args()

    src = Path(args.src)
    out_yml = Path(args.out); out_yml.parent.mkdir(parents=True, exist_ok=True)
    out_map = Path(args.map); out_map.parent.mkdir(parents=True, exist_ok=True)

    if not src.exists():
        print(f"[error] Source folder not found: {src}")
        return

    all_docs = list(src.rglob("*.docx"))
    print(f"[info] Searching in: {src}")
    print(f"[info] Found {len(all_docs)} .docx files")

    claims = []
    claim_map = {}
    counter = defaultdict(int)

    for docx in all_docs:
        lines = extract_from_docx(docx, verbose=args.verbose)
        for txt, kind in lines:
            lhs, rhs = split_identity(txt)
            if lhs is None:
                if args.verbose:
                    print(f"[skip-nonidentity] {docx} :: {txt}")
                continue
            base = docx.stem.replace(" ", "_")
            counter[base] += 1
            cid = f"{base}_{counter[base]:03d}"
            symbols = guess_symbols(lhs, rhs)
            claims.append({
                "id": cid,
                "type": "identity",
                "lhs": lhs,
                "rhs": rhs,
                "symbols": symbols,
                "numeric_trials": 6
            })
            claim_map[cid] = str(docx)
            if args.verbose:
                print(f"[claim] {cid} :: {lhs} = {rhs}")

    if not claims:
        print("[warn] No identity-style equations found.")
    else:
        def yq(s): return '"' + s.replace('"','\\"') + '"'
        with open(out_yml, "w", encoding="utf-8") as f:
            f.write("claims:\n")
            for c in claims:
                f.write(f"  - id: {c['id']}\n")
                f.write(  "    type: identity\n")
                f.write(f"    lhs: {yq(c['lhs'])}\n")
                f.write(f"    rhs: {yq(c['rhs'])}\n")
                if c["symbols"]:
                    f.write(f"    symbols: [{', '.join(c['symbols'])}]\n")
                f.write(f"    numeric_trials: {c['numeric_trials']}\n\n")
        with open(out_map, "w", encoding="utf-8") as f:
            json.dump(claim_map, f, indent=2)
        print(f"[ok] Wrote {out_yml} with {len(claims)} claims")
        print(f"[ok] Wrote {out_map} (claim -> source doc)")

if __name__ == "__main__":
    main()
